"""
app.features.reports.service — Report generation and export (DOCX, PDF, HTML).

Preserved from original with Streamlit-compatible export paths.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.features.clinician_digest.service import ClinicianDigestService
from app.features.companion.service import CompanionService
from app.features.meals.service import MealService
from app.services.repositories import ReportRepository


@dataclass
class ReportRequest:
    report_type: str
    window_days: int = 7


class ReportService:
    def __init__(self, db: Session, patient_id: int):
        self.repo = ReportRepository(db, patient_id)
        self.companion = CompanionService(db, patient_id)
        self.meals = MealService(db, patient_id)
        self.digest = ClinicianDigestService(db, patient_id)
        self.export_dir = get_settings().export_path

    def _compose(self, payload: ReportRequest) -> tuple[str, str]:
        if payload.report_type == "patient_weekly":
            today = self.companion.today_view()
            meal = self.meals.weekly_summary()
            title = "患者周报"
            markdown = "\n".join([
                "# 患者周报",
                "",
                f"- 最新空腹血糖：{today.get('latest_fasting_glucose') or '暂无'}",
                f"- 最新血压：{today.get('latest_blood_pressure') or '暂无'}",
                f"- 待处理提醒：{today.get('pending_reminders', 0)}",
                f"- 本周饮食高风险标签：{', '.join(meal.get('top_risk_tags', [])) or '暂无明显高风险'}",
                "",
                f"- 陪伴建议：{today.get('coach_message', '')}",
            ])
        elif payload.report_type == "adherence_overview":
            today = self.companion.today_view()
            title = "依从性概览"
            markdown = "\n".join([
                "# 依从性概览",
                "",
                f"- 今日待处理提醒：{today.get('pending_reminders', 0)}",
                "- 建议重点查看是否存在连续跳过服药或长期未记录监测值的情况。",
            ])
        else:
            digest = self.digest.generate(payload.window_days)
            title = "门诊前摘要报告"
            markdown = digest.content_markdown
        return title, markdown

    def generate(self, payload: ReportRequest) -> dict:
        title, markdown = self._compose(payload)
        row = self.repo.add(payload.report_type, payload.window_days, title, markdown)
        return {
            "id": row.id,
            "patient_id": row.patient_id,
            "report_type": row.report_type,
            "window_days": row.window_days,
            "title": row.title,
            "markdown": row.markdown,
            "created_at": row.created_at,
        }

    def list_reports(self) -> list[dict]:
        rows = self.repo.list_recent()
        return [
            {
                "id": r.id, "report_type": r.report_type,
                "title": r.title, "markdown": r.markdown,
                "created_at": r.created_at,
            }
            for r in rows
        ]

    def export_docx(self, report_id: int) -> bytes:
        """Generate DOCX bytes for download."""
        from docx import Document

        row = self.repo.get(report_id)
        if row is None:
            raise ValueError("report not found")

        doc = Document()
        for line in row.markdown.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            else:
                doc.add_paragraph(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_html(self, report_id: int) -> str:
        """Generate HTML string for download."""
        row = self.repo.get(report_id)
        if row is None:
            raise ValueError("report not found")
        return (
            '<html><body><pre style="white-space: pre-wrap; font-family: Arial;">'
            + row.markdown
            + "</pre></body></html>"
        )

    def export_pdf(self, report_id: int) -> bytes:
        """Generate PDF bytes for download."""
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas as pdf_canvas

        row = self.repo.get(report_id)
        if row is None:
            raise ValueError("report not found")

        buffer = io.BytesIO()
        c = pdf_canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 40
        for line in row.markdown.splitlines():
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, line[:95])
            y -= 16
        c.save()
        buffer.seek(0)
        return buffer.getvalue()
